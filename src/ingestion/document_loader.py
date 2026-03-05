"""
Document Loader - Load and extract text from various file formats.

Supports: PDF, DOCX, TXT
"""

from typing import Dict, Any, List
from pathlib import Path
import hashlib
from datetime import datetime

from src.utils.logger import setup_logger
from src.utils.exceptions import AgenticRAGException


class DocumentLoadError(AgenticRAGException):
    """Error during document loading."""
    pass


class Document:
    """
    Document container with text and metadata.
    
    Attributes:
        text: Document text content
        metadata: Document metadata (filename, source, etc.)
        doc_id: Unique document identifier
    
    Example:
        >>> doc = Document(
        ...     text="Document content here",
        ...     metadata={"filename": "test.pdf", "source": "upload"}
        ... )
        >>> print(doc.doc_id)  # MD5 hash of text
    """
    
    def __init__(self, text: str, metadata: Dict[str, Any] = None):
        """
        Initialize Document.
        
        Args:
            text: Document text content
            metadata: Optional metadata dictionary
        """
        self.text = text
        self.metadata = metadata or {}
        self.doc_id = self._generate_doc_id()
        
        # Add default metadata
        if "loaded_at" not in self.metadata:
            self.metadata["loaded_at"] = datetime.now().isoformat()
    
    def _generate_doc_id(self) -> str:
        """Generate unique document ID from text hash."""
        return hashlib.md5(self.text.encode()).hexdigest()
    
    def __len__(self) -> int:
        """Get document length."""
        return len(self.text)
    
    def __repr__(self) -> str:
        """String representation."""
        return f"Document(doc_id={self.doc_id[:8]}..., length={len(self.text)})"


class DocumentLoader:
    """
    Load documents from various file formats.
    
    Supports:
    - PDF (.pdf)
    - Word Documents (.docx)
    - Text files (.txt, .md)
    
    Example:
        >>> loader = DocumentLoader()
        >>> doc = loader.load("path/to/file.pdf")
        >>> print(doc.text[:100])
    """
    
    def __init__(self):
        """Initialize document loader."""
        self.logger = setup_logger("document_loader")
        self.supported_formats = [".pdf", ".docx", ".txt", ".md"]
    
    def load(self, file_path: str) -> Document:
        """
        Load document from file.
        
        Args:
            file_path: Path to document file
        
        Returns:
            Document object with text and metadata
        
        Raises:
            DocumentLoadError: If loading fails
        
        Example:
            >>> loader = DocumentLoader()
            >>> doc = loader.load("document.pdf")
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentLoadError(
                message=f"File not found: {file_path}",
                details={"path": file_path}
            )
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            raise DocumentLoadError(
                message=f"Unsupported file format: {file_ext}",
                details={
                    "format": file_ext,
                    "supported": self.supported_formats
                }
            )
        
        self.logger.info(f"Loading document: {path.name}")
        
        try:
            if file_ext == ".pdf":
                text = self._load_pdf(path)
            elif file_ext == ".docx":
                text = self._load_docx(path)
            else:  # .txt, .md
                text = self._load_text(path)
            
            metadata = self._extract_metadata(path, file_ext)
            
            doc = Document(text=text, metadata=metadata)
            
            self.logger.info(
                f"Loaded document: {path.name} "
                f"({len(doc.text)} chars, doc_id={doc.doc_id[:8]}...)"
            )
            
            return doc
            
        except Exception as e:
            self.logger.error(f"Failed to load {path.name}: {str(e)}")
            raise DocumentLoadError(
                message=f"Failed to load document: {str(e)}",
                details={"path": file_path, "error": str(e)}
            ) from e
    
    def count_pages(self, file_path: str) -> int:
        """
        Count pages in document.
        
        Args:
            file_path: Path to document
        
        Returns:
            Number of pages (estimated for non-PDF)
        
        Example:
            >>> loader = DocumentLoader()
            >>> pages = loader.count_pages("document.pdf")
            >>> print(pages)  # 10
        """
        path = Path(file_path)
        
        if not path.exists():
            self.logger.warning(f"File not found: {file_path}")
            return 1
        
        file_ext = path.suffix.lower()
        
        try:
            if file_ext == ".pdf":
                # PDF: Actual page count
                from pypdf import PdfReader
                reader = PdfReader(path)
                return len(reader.pages)
            
            elif file_ext == ".docx":
                # DOCX: Estimate (paragraphs / 30)
                from docx import Document as DocxDocument
                doc = DocxDocument(path)
                paragraph_count = len(doc.paragraphs)
                estimated_pages = max(1, paragraph_count // 30)
                return estimated_pages
            
            else:  # .txt, .md
                # Text: Estimate (lines / 40)
                with open(path, 'r', encoding='utf-8') as f:
                    line_count = len(f.readlines())
                estimated_pages = max(1, line_count // 40)
                return estimated_pages
        
        except Exception as e:
            self.logger.warning(f"Could not count pages for {path.name}: {e}")
            return 1

    def _load_pdf(self, path: Path) -> str:
        """
        Load text from PDF file.
        
        Uses pypdf for text extraction.
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            raise DocumentLoadError(
                message="pypdf not installed. Run: pip install pypdf",
                details={"required_package": "pypdf"}
            )
        
        text_parts = []
        
        reader = PdfReader(path)
        
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        if not text_parts:
            raise DocumentLoadError(
                message="No text content found in PDF",
                details={"path": str(path)}
            )
        
        return "\n\n".join(text_parts)
    
    def _load_docx(self, path: Path) -> str:
        """
        Load text from DOCX file.
        
        Uses python-docx for text extraction.
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise DocumentLoadError(
                message="python-docx not installed. Run: pip install python-docx",
                details={"required_package": "python-docx"}
            )
        
        doc = DocxDocument(path)
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        if not text_parts:
            raise DocumentLoadError(
                message="No text content found in DOCX",
                details={"path": str(path)}
            )
        
        return "\n\n".join(text_parts)
    
    def _load_text(self, path: Path) -> str:
        """
        Load text from plain text file.
        
        Supports .txt, .md files.
        """
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read()
            
            if not text.strip():
                raise DocumentLoadError(
                    message="File is empty",
                    details={"path": str(path)}
                )
            
            return text
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(path, "r", encoding="latin-1") as f:
                    text = f.read()
                return text
            except Exception as e:
                raise DocumentLoadError(
                    message=f"Failed to decode text file: {str(e)}",
                    details={"path": str(path)}
                )
    
    def _extract_metadata(self, path: Path, file_ext: str) -> Dict[str, Any]:
        """
        Extract metadata from file.
        
        Args:
            path: File path
            file_ext: File extension
        
        Returns:
            Metadata dictionary
        """
        stat = path.stat()
        
        return {
            "filename": path.name,
            "file_path": str(path.absolute()),
            "file_type": file_ext,
            "file_size": stat.st_size,
            "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
            "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
            "source": "file_upload"
        }
    
    def load_batch(self, file_paths: List[str]) -> List[Document]:
        """
        Load multiple documents.
        
        Args:
            file_paths: List of file paths
        
        Returns:
            List of Document objects
        
        Example:
            >>> loader = DocumentLoader()
            >>> docs = loader.load_batch(["file1.pdf", "file2.txt"])
        """
        documents = []
        
        for file_path in file_paths:
            try:
                doc = self.load(file_path)
                documents.append(doc)
            except DocumentLoadError as e:
                self.logger.warning(f"Skipping {file_path}: {e.message}")
                continue
        
        self.logger.info(f"Loaded {len(documents)}/{len(file_paths)} documents")
        
        return documents